from datetime import date
import mysql
import mysql.connector
import docx
from decouple import Config, RepositoryEnv
from docx.shared import Pt
import sys



def docx_find_replace_text(doc, search_text, replace_text):
    paragraphs = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraphs.append(paragraph)
    for p in paragraphs:
        if search_text in p.text:
            inline = p.runs
            started = False
            search_index = 0
			
            found_runs = list()
            found_all = False
            replace_done = False
            for i in range(len(inline)):

                # case 1: found in single run so short circuit the replace
                if search_text in inline[i].text and not started:
                    found_runs.append((i, inline[i].text.find(search_text), len(search_text)))
                    text = inline[i].text.replace(search_text, str(replace_text))
                    inline[i].text = text
                    replace_done = True
                    found_all = True
                    break

                if search_text[search_index] not in inline[i].text and not started:
                    continue

                # case 2: search for partial text, find first run
                if search_text[search_index] in inline[i].text and inline[i].text[-1] in search_text and not started:
                    start_index = inline[i].text.find(search_text[search_index])
                    check_length = len(inline[i].text)
                    for text_index in range(start_index, check_length):
                        if inline[i].text[text_index] != search_text[search_index]:
                            break
                    if search_index == 0:
                        started = True
                    chars_found = check_length - start_index
                    search_index += chars_found
                    found_runs.append((i, start_index, chars_found))
                    if search_index != len(search_text):
                        continue
                    else:
                        found_all = True
                        break

                # case 2: search for partial text, find subsequent run
                if search_text[search_index] in inline[i].text and started and not found_all:
                    chars_found = 0
                    check_length = len(inline[i].text)
                    for text_index in range(0, check_length):
                        if inline[i].text[text_index] == search_text[search_index]:
                            search_index += 1
                            chars_found += 1
                        else:
                            break
                    found_runs.append((i, 0, chars_found))
                    if search_index == len(search_text):
                        found_all = True
                        break

            if found_all and not replace_done:
                for i, item in enumerate(found_runs):
                    index, start, length = [t for t in item]
                    if i == 0:
                        text = inline[index].text.replace(inline[index].text[start:start + length], str(replace_text))
                        inline[index].text = text
                    else:
                        text = inline[index].text.replace(inline[index].text[start:start + length], '')
                        inline[index].text = text




def run_script(today= '' , yesterday= ''):

    DOTENV_FILE = r'C:\Users\Haroon Khan Awan\Documents\report\.env'
    env_config = Config(RepositoryEnv(DOTENV_FILE))

    user = env_config.get('DUSER')
    passw = env_config.get('DPASS')
    dhost = env_config.get('DHOST')
    print('')
    print('Establishing connection . . . ')
    connection = mysql.connector.connect(host=dhost,
                                         database='vf',
                                         auth_plugin='mysql_native_password',
                                         user=user,
                                         password=passw)

    print('')
    if (connection.is_connected()):
        print('Established with ', connection.get_server_info())
    else:
        print('Connection not established')
        

    fd = open(r'C:\Users\Haroon Khan Awan\Documents\dailyreport.sql')
    sqlFile = fd.read()
    fd.close()
    
    # Replace the target string
    if today!='' and yesterday!='':
        sqlFile = sqlFile.replace('sysdate()',r"'{}'".format(today))
        sqlFile = sqlFile.replace('date_sub(curdate(), interval 1 DAY)', r"'{}'".format(yesterday))

    sqlCommands = sqlFile.split(';')
    print('')
    print(len(sqlCommands) - 1)
    print('')
    print('Queries have been Filtered!')
    print('')

    cursor = connection.cursor(buffered=True)


    resresult = []
    for query in sqlCommands:
        cursor.execute(query)
        result = cursor.fetchall() #fetchone()
        resresult.append(result)

    print('Queries Results :')
    print('')
    for i in range(len(resresult)-1):
        print(resresult[i][0][0])
    print('')    
        
    today = date.today()
    d2 = today.strftime("%B %d, %Y")
    print(" {} is the date.".format(d2))

    doc = docx.Document(r"C:\Users\Haroon Khan Awan\Documents\report\Report.docx")




    docx_find_replace_text(doc, 'August 3, 2022', " {}".format(d2))
            
            
    table = doc.tables[2]

    for i in range(21):
        if i <= 3:
            table.cell(i+1 , 1).text = str(resresult[i][0][0]);
            paragraph = table.cell(i+1,1).paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size= Pt(10)
            font.bold = True
            font.name = 'Montserrat SemiBold'
        elif i>=4:
            table.cell(i+2 , 1).text = str(resresult[i][0][0]);
            paragraph = table.cell(i+2,1).paragraphs[0]
            run = paragraph.runs
            font = run[0].font
            font.size= Pt(10)
            font.bold = True
            font.name = 'Montserrat SemiBold'
        else:
            pass

        
    table2 = doc.tables[3] 
    x = 21
    for i in range(1,5):
        table2.cell(i , 1).paragraphs[0].text = str(resresult[x][0][0]);
        paragraph = table2.cell(i,1).paragraphs[0]
        run = paragraph.runs
        font = run[0].font
        font.size= Pt(10)
        font.bold = True
        font.name = 'Montserrat SemiBold'
        x+=1    
        
        
        
    doc.save("Digital Pilot {} Report - Block wise.docx".format(d2))


    if connection.is_connected():
        cursor.close()
        connection.close()
        
        

if __name__ == "__main__":  
    if len(sys.argv) == 3:
        run_script(str(sys.argv[1]),str(sys.argv[2]))
    else:
        run_script()
